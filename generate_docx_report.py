import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_report():
    doc = docx.Document()

    # Define Page Margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styling colors
    PRIMARY = RGBColor(15, 23, 42)      # Navy / Charcoal
    SECONDARY = RGBColor(37, 99, 235)   # Royal Blue
    TEXT_DARK = RGBColor(51, 65, 85)    # Slate Dark
    WHITE = RGBColor(255, 255, 255)

    # Set base Normal Style
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = TEXT_DARK
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = SECONDARY
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = SECONDARY
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = TEXT_DARK
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.color.rgb = PRIMARY
        run_t = p.add_run(text)
        run_t.font.color.rgb = TEXT_DARK

    def set_cell_background(cell, fill_color):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    # ==========================================
    # 1. COVER PAGE / TITLE
    # ==========================================
    add_title("PROJECT REPORT & SYNOPSIS\nVOYAGE — TRAVEL PLANNING PLATFORM")
    add_subtitle("Full-Stack Indian Travel Exploration, Budgeting & Cloud Sync Web Platform")

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(36)
    r = p_meta.add_run("Submitted for Academic & Technical Evaluation\nAcademic Year 2026 – 2027\n\nGitHub Repository: https://github.com/pranaykumarsingh06/Mini-Project\nSupabase Cloud Project ID: awdvvaxglwbejfxnvngu")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = SECONDARY

    doc.add_page_break()

    # ==========================================
    # 2. ABSTRACT
    # ==========================================
    add_h1("1. ABSTRACT")
    p_abs = doc.add_paragraph(
        "The VOYAGE Travel Web Application is an end-to-end digital travel planning, itinerary management, sightseeing discovery, "
        "and flight/hotel booking platform designed specifically for Indian domestic and international tourism. Built on modern web technologies "
        "(Python Flask backend, SQLite3 local database, Supabase PostgreSQL cloud sync, and Leaflet.js / OpenStreetMap mapping engines), "
        "VOYAGE solves the fragmentation in traditional travel platforms by combining interactive maps, custom budgeting in Indian Rupees (₹ INR), "
        "instant PDF E-Ticket generation, and administrative governance into a single responsive application.\n\n"
        "The system incorporates real-time asynchronous synchronization with a Supabase cloud database (`awdvvaxglwbejfxnvngu`), ensuring "
        "that every user login, registration, itinerary addition, and booking reservation is persisted securely across serverless cloud environments "
        "with zero infrastructure expenditure overhead."
    )

    # ==========================================
    # 3. INTRODUCTION & OBJECTIVES
    # ==========================================
    add_h1("2. INTRODUCTION & OBJECTIVES")
    add_h2("2.1 Project Overview")
    doc.add_paragraph(
        "Tourism in India is experiencing rapid growth, yet travelers often struggle with fragmented platforms that separate location discovery, "
        "route calculation, budgeting, and reservation management. VOYAGE consolidates these features into an all-in-one web portal with zero learning curve."
    )
    
    add_h2("2.2 Key Project Objectives")
    add_bullet("1. Comprehensive Indian Location Explorer: ", "Provide detailed sightseeing guides, top 6 attraction spots, weather forecasts, best visiting seasons, luxury hotels, and local food spots for 12 major Indian cities (Taj Mahal Agra, Goa, Jaipur, Kerala, Leh Ladakh, Udaipur, Varanasi, Rishikesh, Amritsar, Darjeeling, Shimla, Manali).")
    add_bullet("2. Custom Budgeting Engine: ", "Allow travelers to set custom budgets in ₹ INR, travel dates, and companion styles (Family, Friends, Couple, Solo).")
    add_bullet("3. Dual-Engine Interactive Mapping: ", "Integrate Leaflet.js with OpenStreetMap tiles and Google Maps satellite embeds for 100% map uptime and route calculation.")
    add_bullet("4. Interactive Reservation & E-Ticket Generator: ", "Enable 1-click modal management for flight/hotel seat selection, meal preferences, and printable E-Ticket generation with PNR and barcodes.")
    add_bullet("5. Cloud Data Synchronization: ", "Persist user logins, registrations, trips, and bookings to a remote Supabase PostgreSQL database.")
    add_bullet("6. Administrative Control Portal: ", "Provide a secure admin dashboard (`/admin`) enforcing restricted email authentication (`pranaykrsingh03@gmail.com`) for user and financial analytics management.")

    # ==========================================
    # 4. LITERATURE SURVEY / SYSTEM COMPARISON
    # ==========================================
    add_h1("3. LITERATURE SURVEY & COMPARATIVE ANALYSIS")
    doc.add_paragraph("A comparative evaluation between existing legacy travel platforms and the proposed VOYAGE application:")

    table_comp = doc.add_table(rows=5, cols=4)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_comp.autofit = False

    headers = ["Feature", "Traditional Platforms", "Generic Travel Blogs", "VOYAGE Platform"]
    for i, h in enumerate(headers):
        cell = table_comp.cell(0, i)
        cell.text = h
        set_cell_background(cell, "0F172A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE

    data_comp = [
        ["Sightseeing & Hotel Discovery", "Separated from maps", "Static articles only", "Integrated 1-Click Interactive Guide"],
        ["Budget Planning in ₹ INR", "Fixed commercial prices", "No budget tools", "Customizable ₹ INR Trip Budgeter"],
        ["Cloud Backend Sync", "Proprietary locked DB", "None", "Real-Time Supabase PostgreSQL Sync"],
        ["Admin Control Panel", "Complex enterprise tools", "None", "Minimalist 1-Click Control Portal"]
    ]

    for row_idx, row_data in enumerate(data_comp, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = table_comp.cell(row_idx, col_idx)
            cell.text = text
            set_cell_background(cell, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==========================================
    # 5. SYSTEM REQUIREMENTS
    # ==========================================
    add_h1("4. SYSTEM REQUIREMENTS & SPECIFICATIONS")
    add_h2("4.1 Software Requirements")
    add_bullet("Operating System: ", "Windows 10/11, macOS, or Linux")
    add_bullet("Backend Runtime: ", "Python 3.8+ (Flask Framework 3.0+)")
    add_bullet("Database Engines: ", "SQLite3 (Local Development) & Supabase PostgreSQL (Cloud Production)")
    add_bullet("Frontend Stack: ", "HTML5, CSS3, JavaScript (ES6+), Leaflet.js v1.9.4")
    add_bullet("Deployment Target: ", "Vercel Cloud Serverless Functions")

    add_h2("4.2 Hardware Requirements")
    add_bullet("Client Processor: ", "Dual-Core 1.6 GHz or higher")
    add_bullet("Client Memory: ", "1 GB RAM minimum (50 MB browser RAM footprint)")
    add_bullet("Network: ", "Active Internet Connection for Map Tile Loading & Cloud Sync")

    # ==========================================
    # 6. SYSTEM DESIGN & DIAGRAMS
    # ==========================================
    add_h1("5. SYSTEM DESIGN & DIAGRAMS")
    doc.add_paragraph("The system design follows standard Object-Oriented and Structured System Analysis principles:")

    add_h2("5.1 Use Case Diagram Summary")
    add_bullet("Traveler Actor: ", "Can explore destinations, calculate routes, register accounts, sign in, plan trips in ₹ INR, and manage reservations.")
    add_bullet("Admin Actor: ", "Can log into `/admin` with `pranaykrsingh03@gmail.com`, create users/trips, edit accounts, delete records, and view total revenue.")
    add_bullet("Supabase Cloud Actor: ", "Asynchronously receives user signup, login events, trip itineraries, and booking payloads.")

    add_h2("5.2 Entity-Relationship (ER) Schema")
    add_bullet("USERS: ", "id (PK), full_name, email (UK), password, country, phone, created_at")
    add_bullet("USER_LOGINS: ", "id (PK), user_name, email, login_time, status")
    add_bullet("TRIPS: ", "id (PK), user_name, destination, budget (₹ INR), travel_dates, companion, status, created_at")
    add_bullet("BOOKINGS: ", "id (PK), user_name, destination, booking_type, amount (₹ INR), status, created_at")

    add_h2("5.3 Data Flow Architecture (DFD)")
    add_bullet("DFD Level 0 (Context): ", "Exchanges credentials, trip parameters, and reservation requests between User, VOYAGE Application Core, and Supabase Cloud.")
    add_bullet("DFD Level 1 (Process Decomposition): ", "Sub-processes: 1.0 Auth & Registration, 2.0 Sightseeing Explorer, 3.0 Trip Budgeting, 4.0 Reservation Engine, 5.0 Admin Control Panel.")

    # ==========================================
    # 7. MODULE DESCRIPTION
    # ==========================================
    add_h1("6. MODULE DESCRIPTION")
    add_h2("Module 1: User Authentication & Security")
    doc.add_paragraph("Handles registration, login, session guards, and SHA-256 password hashing to protect user credentials.")

    add_h2("Module 2: Destination & Sightseeing Explorer")
    doc.add_paragraph("Provides comprehensive coverage of 12 Indian tourist destinations with weather, luxury hotels, and sightseeing spots.")

    add_h2("Module 3: Trip Planner & Budgeting")
    doc.add_paragraph("Allows users to configure itineraries, select companion travel styles, and track trip budgets in ₹ INR.")

    add_h2("Module 4: Booking Engine & E-Ticket Generator")
    doc.add_paragraph("Facilitates flight/hotel seat selection, meal customization, reservation status updates, and PDF E-Ticket downloads.")

    add_h2("Module 5: Supabase Cloud Synchronization Engine")
    doc.add_paragraph("Integrates Supabase REST API (`awdvvaxglwbejfxnvngu`) to persist logins, user signups, trips, and bookings to cloud PostgreSQL.")

    add_h2("Module 6: Administrative Control Portal (`/admin`)")
    doc.add_paragraph("Provides executive governance, revenue statistics, user management, and trip record manipulation for project administrators.")

    # ==========================================
    # 8. FEASIBILITY STUDY SUMMARY
    # ==========================================
    add_h1("7. FEASIBILITY STUDY SUMMARY")
    add_bullet("Technical Feasibility: ", "HIGH (10/10) — Built on robust, tested open-source frameworks (Flask, Leaflet, Supabase).")
    add_bullet("Operational Feasibility: ", "HIGH (10/10) — Responsive interface, zero user learning curve, automated cloud logging.")
    add_bullet("Economic Feasibility: ", "HIGH (10/10) — ₹0 Monthly Operating Cost using Vercel, Supabase free-tier, and OpenStreetMap.")
    add_bullet("Legal & Security Feasibility: ", "HIGH (10/10) — Protected with SHA-256 encryption, session guards, and MIT/ODbL licensing compliance.")

    # ==========================================
    # 9. CONCLUSION & FUTURE SCOPE
    # ==========================================
    add_h1("8. CONCLUSION & FUTURE SCOPE")
    doc.add_paragraph(
        "The VOYAGE project successfully demonstrates a scalable, cloud-connected travel exploration and budgeting platform. "
        "By integrating dual-engine interactive maps, custom budgeting in ₹ INR, and automated Supabase cloud sync, VOYAGE provides "
        "a reliable technical solution for tourism platforms."
    )
    add_h2("8.1 Future Scope")
    add_bullet("1. AI Itinerary Generation: ", "Integration of Large Language Models for automated day-wise tour scheduling.")
    add_bullet("2. Payment Gateway Integration: ", "Incorporation of Razorpay / UPI for real-time booking payments.")
    add_bullet("3. Progressive Web App (PWA): ", "Offline caching for mobile devices during mountain/remote region travel.")

    # ==========================================
    # 10. REFERENCES
    # ==========================================
    add_h1("9. REFERENCES")
    add_bullet("1. Flask Documentation: ", "https://flask.palletsprojects.com/")
    add_bullet("2. Supabase Cloud Documentation: ", "https://supabase.com/docs")
    add_bullet("3. Leaflet.js Interactive Maps API: ", "https://leafletjs.com/")
    add_bullet("4. Vercel Serverless Platform: ", "https://vercel.com/docs")
    add_bullet("5. OpenStreetMap Foundation: ", "https://www.openstreetmap.org/")

    # Save document
    filename = "VOYAGE_Project_Report.docx"
    doc.save(filename)
    print(f"Project Report Word document successfully created: {os.path.abspath(filename)}")

if __name__ == '__main__':
    create_report()
